"""Reading documents, judging evidence, and stopping before a change.

Three phases of a real founder objective had no executable capability at
all: understand these documents, compare them, decide what suits. The
Planner was right to refuse -- it is forbidden to invent a capability --
so the gap was never guidance, it was the catalogue.

These tests cover the primitives that close it, and the two properties
that matter more than any of them: **nothing private leaves the machine
without being asked**, and **an original is never changed without a
founder's yes**.

Every document here is synthetic. The founder's own files are never read,
copied, or quoted by a test.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from master_agent.executor.executor import LocalExecutor
from master_agent.permissions.permission_system import (
    ApprovalRequired,
    GrantScope,
    PermissionSystem,
)
from master_agent.plugins.base import RiskTier
from master_agent.plugins.document_plugin import DocumentPlugin
from master_agent.plugins.reasoning_plugin import ReasoningPlugin

ALPHA = "Alpha Consulting. Ran a support desk of four people for two years."
BETA = "Beta Industries. Led three departments totalling sixty people for eight years."


@pytest.fixture
def workspace(tmp_path):
    return {"desktop": tmp_path}


@pytest.fixture
def documents(workspace):
    return DocumentPlugin(LocalExecutor(PermissionSystem()), workspace)


class Recorder:
    """A runner in the shape of `TieredPromptRunner`: one `run(prompt,
    request)`, recording what it was asked and what it was told about the
    work."""

    def __init__(self, answer: str = "a reasoned answer") -> None:
        self.prompts: list[str] = []
        self.requests: list = []
        self._answer = answer

    def run(self, prompt, request, **kwargs):
        self.prompts.append(prompt)
        self.requests.append(request)
        answer = self._answer

        class Outcome:
            ok = True
            text = answer

        return Outcome()


def write_docx(target: Path, lines) -> None:
    from docx import Document as _Docx

    document = _Docx()
    for line in lines:
        document.add_paragraph(line)
    document.save(str(target))


def write_pdf(target: Path, text: str) -> None:
    """A minimal one-page PDF, hand-built so the test owns its fixture and
    does not depend on a writer library it is not testing."""
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET"
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = "%PDF-1.4\n"
    offsets = []
    for index, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n{body}\nendobj\n"
    start = len(out)
    out += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n"
    out += "".join(f"{offset:010d} 00000 n \n" for offset in offsets)
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{start}\n%%EOF\n"
    )
    target.write_bytes(out.encode("latin-1"))


class TestDocumentExtraction:

    def test_a_text_file_is_read(self, documents, workspace):
        (workspace["desktop"] / "notes.txt").write_text(ALPHA, encoding="utf-8")
        result = documents.invoke("extract_text", {"path": "notes.txt"})

        assert result.success, result.error
        assert result.output["text"].strip() == ALPHA
        assert result.output["format"] == "txt"
        assert result.output["filename"] == "notes.txt"

    def test_a_markdown_file_is_read(self, documents, workspace):
        (workspace["desktop"] / "notes.md").write_text(f"# Title\n\n{BETA}", encoding="utf-8")
        result = documents.invoke("extract_text", {"path": "notes.md"})

        assert result.success, result.error
        assert BETA in result.output["text"]
        assert result.output["format"] == "md"

    def test_a_word_document_is_read(self, documents, workspace):
        write_docx(workspace["desktop"] / "profile.docx", [ALPHA, "", BETA])
        result = documents.invoke("extract_text", {"path": "profile.docx"})

        assert result.success, result.error
        assert ALPHA in result.output["text"]
        assert BETA in result.output["text"]
        assert result.output["format"] == "docx"

    def test_a_word_documents_tables_are_read(self, documents, workspace):
        """Real documents put load-bearing content in tables. A reader that
        dropped them would hand the next step a confident, incomplete
        observation."""
        from docx import Document as _Docx

        document = _Docx()
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Languages"
        table.rows[0].cells[1].text = "Marathi, English"
        document.save(str(workspace["desktop"] / "grid.docx"))

        result = documents.invoke("extract_text", {"path": "grid.docx"})
        assert result.success, result.error
        assert "Languages" in result.output["text"]
        assert "Marathi, English" in result.output["text"]

    def test_a_pdf_is_read(self, documents, workspace):
        write_pdf(workspace["desktop"] / "report.pdf", "Quarterly summary text")
        result = documents.invoke("extract_text", {"path": "report.pdf"})

        assert result.success, result.error
        assert "Quarterly summary" in result.output["text"]
        assert result.output["format"] == "pdf"

    def test_an_unsupported_format_is_refused_by_name(self, documents, workspace):
        (workspace["desktop"] / "archive.zip").write_bytes(b"PK\x03\x04")
        result = documents.invoke("extract_text", {"path": "archive.zip"})

        assert not result.success
        assert "unsupported document format" in (result.error or "")

    def test_extraction_never_modifies_the_source(self, documents, workspace):
        target = workspace["desktop"] / "immutable.txt"
        target.write_text(ALPHA, encoding="utf-8")
        before = (target.read_bytes(), target.stat().st_size)

        assert documents.invoke("extract_text", {"path": "immutable.txt"}).success
        assert (target.read_bytes(), target.stat().st_size) == before

    def test_several_documents_can_be_extracted_in_one_step(
        self, documents, workspace
    ):
        """A discovery step returns however many files it found, and a plan
        written before it ran cannot know how many. Without this, "read the
        documents you just found" cannot be planned at all."""
        (workspace["desktop"] / "one.txt").write_text(ALPHA, encoding="utf-8")
        write_docx(workspace["desktop"] / "two.docx", [BETA])

        result = documents.invoke(
            "extract_text", {"path": ["one.txt", "two.docx"]}
        )

        assert result.success, result.error
        assert result.output["documents"] == 2
        assert ALPHA in result.output["text"]
        assert BETA in result.output["text"]

    def test_each_document_is_named_above_its_own_text(self, documents, workspace):
        """An undivided blob would let a later step attribute one
        document's claim to another."""
        (workspace["desktop"] / "one.txt").write_text(ALPHA, encoding="utf-8")
        (workspace["desktop"] / "two.txt").write_text(BETA, encoding="utf-8")

        text = documents.invoke(
            "extract_text", {"path": ["one.txt", "two.txt"]}
        ).output["text"]

        assert "--- one.txt ---" in text
        assert "--- two.txt ---" in text
        assert text.index("one.txt") < text.index(ALPHA)

    def test_a_document_that_could_not_be_read_is_reported(
        self, documents, workspace
    ):
        """A comparison made over two of three documents while claiming
        three is a wrong answer."""
        (workspace["desktop"] / "one.txt").write_text(ALPHA, encoding="utf-8")

        result = documents.invoke(
            "extract_text", {"path": ["one.txt", "absent.txt"]}
        )

        assert result.success
        assert result.output["documents"] == 1
        assert len(result.output["unreadable"]) == 1

    def test_the_single_document_shape_is_unchanged(self, documents, workspace):
        """Every existing caller and binding must keep working."""
        (workspace["desktop"] / "solo.txt").write_text(ALPHA, encoding="utf-8")
        output = documents.invoke("extract_text", {"path": "solo.txt"}).output

        assert output["filename"] == "solo.txt"
        assert output["text"].strip() == ALPHA
        assert output["documents"] == 1
        assert "--- " not in output["text"]

    def test_extraction_is_read_only(self, documents):
        tiers = {c.name: c.risk_tier for c in documents.manifest.capabilities}
        assert tiers["extract_text"] is RiskTier.READ_ONLY


class TestReasoningOverEvidence:

    def test_it_reasons_over_text_supplied_by_an_earlier_step(self, workspace):
        runner = Recorder("Beta Industries has the wider scope.")
        reasoning = ReasoningPlugin(LocalExecutor(PermissionSystem()), runner)

        result = reasoning.invoke("transform", {
            "instruction": "Which of these has more leadership scope?",
            "context": f"{ALPHA}\n\n{BETA}",
        })

        assert result.success, result.error
        assert result.output["text"] == "Beta Industries has the wider scope."
        # The evidence reached the provider rather than being summarised
        # away by the capability itself.
        assert BETA in runner.prompts[0]

    def test_several_documents_can_feed_one_reasoning_step(self, documents, workspace):
        """The comparison case: two sources extracted, both reasoned over."""
        (workspace["desktop"] / "one.txt").write_text(ALPHA, encoding="utf-8")
        write_docx(workspace["desktop"] / "two.docx", [BETA])

        first = documents.invoke("extract_text", {"path": "one.txt"}).output["text"]
        second = documents.invoke("extract_text", {"path": "two.docx"}).output["text"]

        runner = Recorder("Two sources compared.")
        reasoning = ReasoningPlugin(LocalExecutor(PermissionSystem()), runner)
        result = reasoning.invoke("transform", {
            "instruction": "Compare these documents.",
            "context": f"{first}\n\n---\n\n{second}",
        })

        assert result.success
        assert ALPHA in runner.prompts[0]
        assert BETA in runner.prompts[0]

    def test_its_result_can_feed_a_later_step(self, workspace):
        """The output is a plain string, which is what makes it bindable."""
        runner = Recorder("senior operations leadership, Pune")
        reasoning = ReasoningPlugin(LocalExecutor(PermissionSystem()), runner)
        produced = reasoning.invoke("transform", {"instruction": "Derive search terms."})

        documents = DocumentPlugin(LocalExecutor(PermissionSystem()), workspace)
        written = documents.invoke("write_document", {
            "path": "criteria.txt", "content": produced.output["text"],
        })

        assert written.success, written.error
        assert (workspace["desktop"] / "criteria.txt").read_text(
            encoding="utf-8"
        ) == "senior operations leadership, Pune"

    def test_required_sections_are_reported_when_absent(self, workspace):
        """A mechanical check, and only that: it says the word is missing,
        never that the judgement is wrong."""
        reasoning = ReasoningPlugin(
            LocalExecutor(PermissionSystem()), Recorder("strengths only")
        )
        result = reasoning.invoke("transform", {
            "instruction": "Summarise.", "must_contain": ["strengths", "gaps"],
        })

        assert result.success
        assert result.output["missing_sections"] == ["gaps"]

    def test_an_empty_answer_is_a_failure(self, workspace):
        """A later step binding to emptiness would carry it forward
        silently."""
        reasoning = ReasoningPlugin(LocalExecutor(PermissionSystem()), Recorder("   "))
        result = reasoning.invoke("transform", {"instruction": "Summarise."})

        assert not result.success
        assert "empty" in (result.error or "").lower()

    def test_reasoning_is_read_only(self, workspace):
        reasoning = ReasoningPlugin(LocalExecutor(PermissionSystem()), Recorder())
        tiers = {c.name: c.risk_tier for c in reasoning.manifest.capabilities}
        assert tiers["transform"] is RiskTier.READ_ONLY

    def test_reasoning_cannot_touch_the_environment(self, workspace):
        """Brain decides, hands act. The capability publishes exactly one
        output and no way to cause an effect."""
        reasoning = ReasoningPlugin(LocalExecutor(PermissionSystem()), Recorder())
        capability = reasoning.manifest.capabilities[0]
        outputs = {
            field["name"]
            for field in (capability.output_schema or {}).get("fields", [])
        }

        assert "text" in outputs
        for forbidden in ("path", "url", "session_id", "written", "sent"):
            assert forbidden not in outputs


class TestPrivacy:
    """The rule: nothing private leaves the machine without being asked."""

    def test_evidence_is_treated_as_private_by_default(self, workspace):
        runner = Recorder()
        reasoning = ReasoningPlugin(LocalExecutor(PermissionSystem()), runner)
        reasoning.invoke("transform", {
            "instruction": "Summarise this.", "context": "a personal document",
        })

        assert runner.requests[0].sensitive is True, (
            "document evidence was routed as unrestricted by default"
        )

    def test_a_plan_may_declare_material_public(self, workspace):
        runner = Recorder()
        reasoning = ReasoningPlugin(LocalExecutor(PermissionSystem()), runner)
        reasoning.invoke("transform", {
            "instruction": "Rank these.", "context": "a public web page",
            "sensitive": False,
        })

        assert runner.requests[0].sensitive is False

    def test_a_private_request_to_a_public_provider_needs_the_founder(self):
        """The existing rule, asserted where it now matters: sensitive work
        may not reach a non-PRIVATE provider without an approval."""
        from dataclasses import dataclass

        from master_agent.ai_infrastructure.approval import (
            SENSITIVE_THIRD_PARTY,
            approval_needed,
        )
        from master_agent.ai_infrastructure.catalog import PROVIDER_CATALOG

        @dataclass
        class Profile:
            """The two facts the rule is about. `approval_needed` reads a
            resolved profile's `cost`, not a catalogue spec's
            `cost_per_call` -- passing the spec makes every provider look
            paid and the privacy assertion vacuous."""

            cost: float
            privacy: str

        by_id = {spec.provider_id: spec for spec in PROVIDER_CATALOG}
        local_spec = by_id["ollama.local"]
        cloud_spec = by_id["gemini.api"]

        # The catalogue facts this rule depends on, asserted rather than
        # assumed: local is free and private, Gemini is not private.
        assert local_spec.cost_per_call == 0.0
        assert local_spec.privacy == "private"
        assert cloud_spec.privacy != "private"

        local = Profile(cost=local_spec.cost_per_call, privacy=local_spec.privacy)
        cloud = Profile(cost=cloud_spec.cost_per_call, privacy=cloud_spec.privacy)

        assert approval_needed(local, "sensitive") is None, (
            "a local private provider should need no approval"
        )
        assert approval_needed(cloud, "sensitive") == SENSITIVE_THIRD_PARTY

    def _ladder(self):
        from master_agent.ai_infrastructure.tiered_runner import TieredPromptRunner

        return TieredPromptRunner(
            object(),
            local_provider_ids=frozenset({"ollama.local"}),
            gemini_provider_ids=frozenset({"gemini.api"}),
            desktop_provider_ids=frozenset({"chatgpt-desktop"}),
            browser_provider_ids=frozenset({"browser.free-ai"}),
        )

    def test_sensitive_work_goes_to_the_machines_own_runtime_first(self):
        from master_agent.ai_infrastructure.tiered_runner import TIER_LOCAL

        class Request:
            sensitive = True

        order = [name for name, _ in self._ladder()._ordered_tiers(Request())]
        assert order[0] == TIER_LOCAL

    def test_public_work_keeps_the_original_ladder(self):
        """Order is capability, not privacy -- the Broker already refuses to
        send sensitive work anywhere non-private. Forcing every request
        through the local runtime would buy no privacy and cost a quarter
        of an hour per planning call."""
        class Request:
            sensitive = False

        order = [name for name, _ in self._ladder()._ordered_tiers(Request())]
        assert order[:3] == ["gemini", "desktop", "browser"]
        assert order[-1] == "local", "the local runtime should remain a fallback"

    def test_a_request_that_says_nothing_is_treated_as_public(self):
        """`sensitive` is absent on some request objects; the ladder must
        not crash, and reasoning defaults to sensitive at its own layer."""
        order = [name for name, _ in self._ladder()._ordered_tiers(object())]
        assert order[0] == "gemini"

    def test_existing_ladders_are_unchanged_without_a_local_tier(self):
        """Additive: a caller that names no local provider keeps exactly
        the ladder it had."""
        from master_agent.ai_infrastructure.tiered_runner import TieredPromptRunner

        runner = TieredPromptRunner(
            object(),
            gemini_provider_ids=frozenset({"gemini.api"}),
            desktop_provider_ids=frozenset({"chatgpt-desktop"}),
            browser_provider_ids=frozenset({"browser.free-ai"}),
        )
        populated = [name for name, ids in runner._tiers if ids]
        assert populated == ["gemini", "desktop", "browser"]


class TestChangingSomethingRequiresAYes:

    def test_reading_and_reasoning_need_no_approval(self, workspace):
        """Discovery is harmless. A mission that asked permission to read
        would teach a founder to click yes without looking."""
        permissions = PermissionSystem()
        permissions.check("document", "extract_text", RiskTier.READ_ONLY)
        permissions.check("reasoning", "transform", RiskTier.READ_ONLY)

    def test_writing_a_document_stops_for_the_founder(self):
        permissions = PermissionSystem()
        with pytest.raises(ApprovalRequired):
            permissions.check("document", "write_document", RiskTier.REVERSIBLE_WRITE)

    def test_the_same_work_continues_once_the_founder_says_yes(self, workspace):
        """Approval resumes the work rather than restarting it."""
        permissions = PermissionSystem()
        executor = LocalExecutor(permissions)
        documents = DocumentPlugin(executor, workspace)

        with pytest.raises(ApprovalRequired):
            permissions.check("document", "write_document", RiskTier.REVERSIBLE_WRITE)

        permissions.grant("document", "write_document", GrantScope.ONCE)
        permissions.check("document", "write_document", RiskTier.REVERSIBLE_WRITE)

        result = documents.invoke("write_document", {
            "path": "revised.txt", "content": "the approved revision",
        })
        assert result.success, result.error
        assert (workspace["desktop"] / "revised.txt").exists()

    def test_writing_is_gated_by_risk_tier(self, documents):
        tiers = {c.name: c.risk_tier for c in documents.manifest.capabilities}
        assert tiers["write_document"] is RiskTier.REVERSIBLE_WRITE


class TestOriginalsSurvive:

    def test_a_revision_is_a_new_file_and_the_original_is_untouched(
        self, documents, workspace
    ):
        original = workspace["desktop"] / "profile.docx"
        write_docx(original, [ALPHA])
        before = original.read_bytes()

        result = documents.invoke("write_document", {
            "path": "profile_revised.docx",
            "content": f"{ALPHA}\nAdded a summary line.",
            "format": "docx",
        })

        assert result.success, result.error
        assert (workspace["desktop"] / "profile_revised.docx").exists()
        assert original.read_bytes() == before, "the original was modified"

    def test_an_existing_file_is_not_replaced_by_accident(self, documents, workspace):
        target = workspace["desktop"] / "keep.txt"
        target.write_text("original", encoding="utf-8")

        result = documents.invoke("write_document", {
            "path": "keep.txt", "content": "replacement",
        })

        assert not result.success
        assert "already exists" in (result.error or "")
        assert target.read_text(encoding="utf-8") == "original"

    def test_a_docx_is_a_real_word_document(self, documents, workspace):
        """Text under a .docx name is a file Word refuses to open, and a
        mission that reported success while writing one would be lying in
        the most practical possible way."""
        from docx import Document as _Docx

        assert documents.invoke("write_document", {
            "path": "written.docx", "content": "First line\nSecond line",
            "format": "docx",
        }).success

        reopened = _Docx(str(workspace["desktop"] / "written.docx"))
        assert [p.text for p in reopened.paragraphs][:2] == ["First line", "Second line"]

    def test_a_format_may_not_disagree_with_the_extension(self, documents):
        result = documents.invoke("write_document", {
            "path": "pretend.docx", "content": "plain text", "format": "txt",
        })
        assert not result.success
        assert "must be what its name says it is" in (result.error or "")

    def test_the_written_document_can_be_read_back(self, documents, workspace):
        """Verification of a change means observing it, not trusting the
        step that made it."""
        documents.invoke("write_document", {
            "path": "roundtrip.docx", "content": "Verified content", "format": "docx",
        })
        result = documents.invoke("extract_text", {"path": "roundtrip.docx"})

        assert result.success, result.error
        assert "Verified content" in result.output["text"]


class TestThePlaybookIsGeneric:

    def test_it_covers_multi_step_and_phased_work(self):
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "multi-step work" in text
        assert "phases" in text
        assert "bounded" in text

    def test_it_teaches_evidence_before_judgement(self):
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "evidence before judgement" in text
        assert "guess wearing a result's clothes" in text

    def test_it_teaches_inspecting_several_sources(self):
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "several sources" in text
        assert "accident that looks like one" in text

    def test_it_teaches_that_ordinary_work_runs_itself(self):
        """This test used to assert the opposite, and was wrong. It pinned
        "approval belongs immediately before the first step that actually
        changes something" -- which is not Founder Edition's policy. A
        reversible write is pre-granted and runs automatically; only
        destructive, financial and privacy actions are held."""
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "ordinary work runs on its own" in text
        assert "none of it waits for anyone" in text
        assert "that is not the plan's job" in text

    def test_it_does_not_teach_approval_before_every_change(self):
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "approval belongs" not in text
        assert "get approval if the next step changes something" not in text

    def test_it_keeps_the_founder_requested_exception(self):
        """Asking to see something first is part of the objective, not a
        policy matter."""
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "what the founder asked for out loud" in text
        assert "part of the outcome, not a policy matter" in text

    def test_it_teaches_following_the_objective_about_originals(self):
        """This asserted a universal "never overwrite" rule, which came
        from one objective that happened to say so. Kalpavriksha follows
        what was asked: preserve when a copy was requested, replace when
        replacement was requested, and never infer replacement from
        silence."""
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "do not decide it for them" in text
        assert "leave the original where it is" in text
        assert "not a detail to infer" in text

    def test_it_requires_the_answer_to_be_delivered_somewhere(self):
        """A step that works the answer out and hands it to nobody has not
        delivered it."""
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "somewhere they can actually open" in text
        assert "hands it to nobody has not delivered it" in text

    def test_it_forbids_unobserved_external_claims(self):
        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert "must come from something the mission actually observed" in text
        assert "not stated" in text

    def test_the_playbook_reaches_the_planning_prompt(self):
        from master_agent.planner.catalogue import CapabilityOption
        from master_agent.planner.plan import Intent
        from master_agent.planner.prompting import build_prompt

        prompt = build_prompt(
            Intent(goal="anything at all"),
            (CapabilityOption(name="Filesystem.ReadFile", required_args=("path",),
                              args_complete=True),),
        )
        assert "Evidence before judgement" in prompt

    @pytest.mark.parametrize("word", [
        "cv", "resume", "curriculum", "vacancy", "career", "recruit",
        "linkedin", "naukri", "indeed", "d drive", "salary",
    ])
    def test_no_task_specific_vocabulary_appears(self, word):
        """Written after one founder objective. If it names that objective,
        it only works for that objective."""
        import re

        from master_agent.planner.task_playbook import playbook_lines

        text = " ".join(playbook_lines()).lower()
        assert not re.search(rf"(?<![a-z]){re.escape(word)}(?![a-z])", text), (
            f"the playbook names {word!r}"
        )


class TestNoTaskSpecificBranching:
    """The capabilities were built for one founder objective and must not
    know it existed."""

    @pytest.mark.parametrize("module", [
        "master_agent.executor.actions.document.extract_text",
        "master_agent.executor.actions.document.write_document",
        "master_agent.executor.actions.reasoning.transform",
        "master_agent.executor.actions.browser.read_page_text",
        "master_agent.planner.task_playbook",
    ])
    def test_no_module_branches_on_task_vocabulary(self, module):
        import importlib
        import inspect
        import re

        source = inspect.getsource(importlib.import_module(module)).lower()
        # Branching, specifically: the word may appear in prose explaining
        # why the module exists, but never in a condition.
        for word in ("cv", "resume", "curriculum", "job", "vacancy", "career"):
            assert not re.search(rf"if\b[^\n]*['\"]{word}['\"]", source), (
                f"{module} branches on {word!r}"
            )
